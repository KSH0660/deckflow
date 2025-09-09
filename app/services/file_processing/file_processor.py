import io
from pathlib import Path

from docx import Document
from PIL import Image
from PyPDF2 import PdfReader

from app.logging import get_logger

logger = get_logger(__name__)


class FileProcessor:
    """파일 타입별 텍스트 추출 및 처리 클래스"""

    ALLOWED_EXTENSIONS = {
        ".txt",
        ".md",
        ".pdf",
        ".docx",
        ".doc",
        ".jpg",
        ".jpeg",
        ".png",
        ".gif",
        ".bmp",
    }

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @classmethod
    def is_allowed_file(cls, filename: str) -> bool:
        """허용된 파일 확장자인지 확인"""
        return Path(filename).suffix.lower() in cls.ALLOWED_EXTENSIONS

    @classmethod
    def is_valid_size(cls, file_size: int) -> bool:
        """파일 크기가 허용 범위인지 확인"""
        return file_size <= cls.MAX_FILE_SIZE

    @classmethod
    async def extract_text(cls, file_content: bytes, filename: str) -> str:
        """파일 내용에서 텍스트 추출"""
        extension = Path(filename).suffix.lower()
        file_size_kb = len(file_content) / 1024

        logger.info(
            "📄 [FILE_PROCESSOR] 파일 텍스트 추출 시작",
            filename=filename,
            extension=extension,
            file_size_kb=round(file_size_kb, 2),
        )

        try:
            if extension == ".txt":
                extracted_text = cls._extract_from_txt(file_content)
                parser_library = "python built-in (decode)"
            elif extension == ".md":
                extracted_text = cls._extract_from_txt(file_content)
                parser_library = "python built-in (decode)"
            elif extension == ".pdf":
                extracted_text = cls._extract_from_pdf(file_content)
                parser_library = "PyPDF2"
            elif extension in [".docx", ".doc"]:
                extracted_text = cls._extract_from_docx(file_content)
                parser_library = "python-docx"
            elif extension in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
                extracted_text = cls._extract_from_image(file_content, filename)
                parser_library = "Pillow (PIL)"
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {extension}")

            text_length = len(extracted_text)
            logger.info(
                "✅ [FILE_PROCESSOR] 파일 텍스트 추출 완료",
                filename=filename,
                extension=extension,
                parser_library=parser_library,
                extracted_text_length=text_length,
                file_size_kb=round(file_size_kb, 2),
                text_preview=extracted_text[:100]
                + ("..." if text_length > 100 else ""),
            )

            # 파일 내용이 길다면 요약 적용
            from app.services.file_processing import summarize_file_content

            logger.info(
                "🔍 [FILE_PROCESSOR] 파일 내용 요약 검사 중",
                filename=filename,
                text_length=text_length,
            )

            processed_text = await summarize_file_content(extracted_text, filename)

            if len(processed_text) != text_length:
                logger.info(
                    "📝 [FILE_PROCESSOR] 파일 내용 요약 적용됨",
                    filename=filename,
                    original_length=text_length,
                    processed_length=len(processed_text),
                    compression_ratio=round(len(processed_text) / text_length, 2),
                )
            else:
                logger.info(
                    "📝 [FILE_PROCESSOR] 파일 내용 요약 불필요 - 원본 사용",
                    filename=filename,
                    text_length=text_length,
                )

            return processed_text

        except Exception as e:
            logger.error(
                "❌ [FILE_PROCESSOR] 파일 텍스트 추출 실패",
                filename=filename,
                extension=extension,
                file_size_kb=round(file_size_kb, 2),
                error=str(e),
            )
            raise

    @staticmethod
    def _extract_from_txt(file_content: bytes) -> str:
        """텍스트 파일에서 내용 추출"""
        try:
            # UTF-8으로 먼저 시도
            text = file_content.decode("utf-8")
            logger.debug("📝 [TXT_PROCESSOR] UTF-8 인코딩으로 성공적으로 디코드")
            return text
        except UnicodeDecodeError:
            logger.debug("📝 [TXT_PROCESSOR] UTF-8 실패, EUC-KR 시도")
            # 실패하면 EUC-KR로 시도
            try:
                text = file_content.decode("euc-kr")
                logger.debug("📝 [TXT_PROCESSOR] EUC-KR 인코딩으로 성공적으로 디코드")
                return text
            except UnicodeDecodeError:
                logger.debug("📝 [TXT_PROCESSOR] EUC-KR 실패, latin-1로 폴백")
                # 마지막으로 latin-1로 시도
                text = file_content.decode("latin-1", errors="ignore")
                logger.warning(
                    "📝 [TXT_PROCESSOR] latin-1 폴백 사용 (일부 문자가 손실될 수 있음)"
                )
                return text

    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            page_count = len(pdf_reader.pages)

            logger.debug("📋 [PDF_PROCESSOR] PDF 파일 분석 시작", page_count=page_count)

            text = ""
            for i, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                text += page_text + "\n"
                logger.debug(
                    f"📋 [PDF_PROCESSOR] 페이지 {i}/{page_count} 처리 완료",
                    page_text_length=len(page_text),
                )

            final_text = text.strip()
            logger.debug(
                "📋 [PDF_PROCESSOR] PDF 텍스트 추출 완료",
                total_pages=page_count,
                total_text_length=len(final_text),
            )
            return final_text
        except Exception as e:
            logger.error("📋 [PDF_PROCESSOR] PDF 처리 중 오류 발생", error=str(e))
            return f"PDF 처리 중 오류 발생: {str(e)}"

    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            paragraph_count = len(doc.paragraphs)

            logger.debug(
                "📄 [DOCX_PROCESSOR] DOCX 파일 분석 시작",
                paragraph_count=paragraph_count,
            )

            text = ""
            non_empty_paragraphs = 0
            for _i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if para_text.strip():  # 비어있지 않은 문단만 카운트
                    non_empty_paragraphs += 1
                text += para_text + "\n"

            final_text = text.strip()
            logger.debug(
                "📄 [DOCX_PROCESSOR] DOCX 텍스트 추출 완료",
                total_paragraphs=paragraph_count,
                non_empty_paragraphs=non_empty_paragraphs,
                total_text_length=len(final_text),
            )
            return final_text
        except Exception as e:
            logger.error("📄 [DOCX_PROCESSOR] DOCX 처리 중 오류 발생", error=str(e))
            return f"DOCX 처리 중 오류 발생: {str(e)}"

    @staticmethod
    def _extract_from_image(file_content: bytes, filename: str) -> str:
        """이미지 파일 처리 (OCR 없이 파일 정보만 반환)"""
        try:
            image = Image.open(io.BytesIO(file_content))
            width, height = image.size
            format_name = image.format or "Unknown"

            logger.debug(
                "🖼️ [IMAGE_PROCESSOR] 이미지 파일 분석 완료",
                filename=filename,
                format=format_name,
                width=width,
                height=height,
                mode=image.mode,
                has_transparency=image.mode in ("RGBA", "LA"),
            )

            info_text = f"""이미지 파일 정보:
파일명: {filename}
형식: {format_name}
크기: {width} x {height}
모드: {image.mode}

참고: 이미지의 텍스트 내용을 추출하려면 OCR 기능이 필요합니다."""
            return info_text
        except Exception as e:
            logger.error(
                "🖼️ [IMAGE_PROCESSOR] 이미지 처리 중 오류 발생",
                filename=filename,
                error=str(e),
            )
            return f"이미지 처리 중 오류 발생: {str(e)}"


class FileStorage:
    """파일 저장소 관리 클래스"""

    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)

    def save_file(self, filename: str, content: bytes) -> str:
        """파일을 저장하고 저장된 경로 반환"""
        # 안전한 파일명 생성
        safe_filename = self._get_safe_filename(filename)
        file_path = self.upload_dir / safe_filename

        # 중복 파일명 처리
        counter = 1
        original_path = file_path
        while file_path.exists():
            name_part = original_path.stem
            ext_part = original_path.suffix
            file_path = self.upload_dir / f"{name_part}_{counter}{ext_part}"
            counter += 1

        # 파일 저장
        with open(file_path, "wb") as f:
            f.write(content)

        return str(file_path)

    def delete_file(self, file_path: str) -> bool:
        """파일 삭제"""
        try:
            Path(file_path).unlink(missing_ok=True)
            return True
        except Exception:
            return False

    def _get_safe_filename(self, filename: str) -> str:
        """안전한 파일명으로 변환"""
        # 위험한 문자들 제거
        safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._-"
        safe_name = "".join(c if c in safe_chars else "_" for c in filename)

        # 길이 제한
        if len(safe_name) > 255:
            name_part = safe_name[:200]
            ext_part = Path(safe_name).suffix
            safe_name = name_part + ext_part

        return safe_name
